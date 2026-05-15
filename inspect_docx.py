from docx import Document
from docx.oxml.ns import qn
import re

def inspect():
    file_path = '/home/ubuntu/rencipe/docs/CPSC597 Project Report Draft (2).docx'
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"Error opening document: {e}")
        return

    print("Document opens successfully.")

    # 1. Title is updated (Assume first paragraph or some logic)
    # Let's look at the first few paragraphs
    title_updated = False
    for p in doc.paragraphs[:5]:
        if p.text.strip():
            print(f"Title candidate: {p.text}")
            title_updated = True
            break
    
    # 2. Old markers absent
    content = ""
    for p in doc.paragraphs:
        content += p.text + "\n"
    
    markers = ["(incomplete)", "NOT STARTED"]
    for m in markers:
        count = content.count(m)
        print(f"Marker '{m}' count: {count}")

    # 3. Not all non-empty paragraphs are blue
    # Standard blue colors in Word: e.g., '0000FF', '1F497D', '4F81BD'
    blue_count = 0
    total_non_empty = 0
    for p in doc.paragraphs:
        if p.text.strip():
            total_non_empty += 1
            is_blue = False
            for run in p.runs:
                if run.font.color and run.font.color.rgb:
                    # Very rough check for blue-ish colors
                    rgb = str(run.font.color.rgb)
                    if rgb.startswith('00') or rgb.endswith('FF') or '4F81BD' in rgb:
                        is_blue = True
                        break
            if is_blue:
                blue_count += 1
    
    blue_percentage = (blue_count / total_non_empty * 100) if total_non_empty > 0 else 0
    print(f"Non-empty paragraphs: {total_non_empty}, Blue: {blue_count} ({blue_percentage:.2f}%)")

    # 4. Added bullet paragraphs use real w:numPr numbering
    num_pr_bullets = 0
    para_list_items = []
    for p in doc.paragraphs:
        if p._element.xpath('./w:pPr/w:numPr'):
            num_pr_bullets += 1
            para_list_items.append(p.text[:30])
    print(f"Paragraphs with w:numPr numbering: {num_pr_bullets}")

    # 5. Body tables after title table are real w:tbl and have explicit tblBorders
    tables = doc.tables
    print(f"Total tables: {len(tables)}")
    for i, table in enumerate(tables):
        # Skip the first table (title table potentially)
        if i == 0: continue
        
        tbl_pr = table._element.xpath('./w:tblPr/w:tblBorders')
        has_borders = len(tbl_pr) > 0
        print(f"Table {i+1} has explicit tblBorders: {has_borders}")

    # 6. Screenshot suggestion lines exist
    # Look for patterns like [Screenshot ...] or Suggestion: Screenshot
    screenshot_suggestions = re.findall(r"(?i)screenshot.*suggestion|suggested.*screenshot|\[screenshot\]", content)
    print(f"Screenshot suggestion hits: {len(screenshot_suggestions)}")

    # 7. Main sections 1 through 9 exist
    sections_found = {}
    for i in range(1, 10):
        # Look for "1. ", "1 " at start of paragraph
        pattern = re.compile(rf"^\s*{i}[\.\s]")
        found = False
        for p in doc.paragraphs:
            if pattern.match(p.text.strip()):
                found = True
                sections_found[i] = p.text.strip()
                break
        if not found:
            sections_found[i] = None
    
    for i in range(1, 10):
        print(f"Section {i}: {'Found (' + sections_found[i][:40] + '...)' if sections_found[i] else 'Not Found'}")

inspect()
