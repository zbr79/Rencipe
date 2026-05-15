from docx import Document
doc = Document("/home/ubuntu/rencipe/docs/CPSC597 Project Report Draft (2).docx")
found_heading_4 = False
for p in doc.paragraphs:
    if "4. System Architecture" in p.text:
        found_heading_4 = True
        print("Found Architecture section...")
        continue
    if found_heading_4:
        # Stop at next Heading 1
        if p.style.name == 'Heading 1' and "5." in p.text:
            print("Reached next section.")
            break
        # Print tables in between
        # This is tricky because tables aren't in paragraphs.
        # But we can see which tables are after which paragraphs in the internal XML.
        # Alternatively, let's just look for any table with architecture-related words.
        pass

print("\n--- All tables search for architecture or system components ---")
for i, table in enumerate(doc.tables):
    text = " ".join(cell.text for row in table.rows for cell in row.cells).lower()
    if any(word in text for word in ["component", "module", "layer", "frontend", "backend", "database", "infrastructure"]):
         print(f"Table {i} might be architecture: {text[:100]}...")
