from docx import Document
doc = Document("/home/ubuntu/rencipe/docs/CPSC597 Project Report Draft (2).docx")

# Searching for Heading 4 (Architecture)
print("--- Checking Headings 4 to 6 ---")
for p in doc.paragraphs:
    if "4" in p.text and "Architecture" in p.text:
        print(f"Found Architecture Heading: {p.text} (Style: {p.style.name})")
    if "5" in p.text and "Implementation" in p.text:
         print(f"Found Implementation Heading: {p.text} (Style: {p.style.name})")
    if "6" in p.text and "Testing" in p.text:
         print(f"Found Testing Heading: {p.text} (Style: {p.style.name})")

# Searching for Tables again with more lenient keywords
print("\n--- Searching Tables for keywords ---")
keywords = ["goal", "requirement", "architecture", "testing", "functional"]
for i, table in enumerate(doc.tables):
    text = " ".join(cell.text for row in table.rows for cell in row.cells).lower()
    for kw in keywords:
        if kw in text:
            print(f"Table {i} contains keyword '{kw}'")
            # print(f"Snippet: {text[:100]}")

