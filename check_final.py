from docx import Document
import re

doc = Document('/home/ubuntu/rencipe/docs/CPSC597 Project Report Draft (2).docx')

# 1. Title is updated
title_table = doc.tables[0]
title_text = title_table.rows[0].cells[0].text.strip()
print(f"Title in table: {title_text}")

# 2. Old markers absent
full_text = "\n".join([p.text for p in doc.paragraphs])
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            full_text += "\n" + cell.text

incomplete_count = full_text.count("(incomplete)")
not_started_count = full_text.count("NOT STARTED")
print(f"Markers: (incomplete)={incomplete_count}, NOT STARTED={not_started_count}")

# 3. Not all non-empty paragraphs are blue
total_p = 0
blue_p = 0
for p in doc.paragraphs:
    if p.text.strip():
        total_p += 1
        is_blue = False
        for run in p.runs:
            if run.font.color and run.font.color.rgb:
                if str(run.font.color.rgb) in ['0000FF', '1F497D', '4F81BD']:
                    is_blue = True
                    break
        if is_blue: blue_p += 1
print(f"Paragraphs: Total={total_p}, Blue={blue_p} ({blue_p/total_p*100:.1f}%)")

# 4. Bullet points
num_pr = 0
for p in doc.paragraphs:
    if p._element.xpath('./w:pPr/w:numPr'):
        num_pr += 1
print(f"Numbering (w:numPr): {num_pr}")

# 5. Tables and borders
print(f"Tables count: {len(doc.tables)}")
for i, t in enumerate(doc.tables[1:], 2):
    borders = t._element.xpath('./w:tblPr/w:tblBorders')
    print(f"Table {i} borders: {'Yes' if borders else 'No'}")

# 6. Screenshot suggestions
screenshot_hits = len(re.findall(r"(?i)screenshot.*suggestion|suggested.*screenshot|\[screenshot\]", full_text))
print(f"Screenshot suggestions: {screenshot_hits}")

# 7. Sections 1-9
for i in range(1, 10):
    pattern = re.compile(rf"^\s*{i}[\.\s]")
    found = any(pattern.match(p.text.strip()) for p in doc.paragraphs)
    print(f"Section {i}: {'Found' if found else 'NOT FOUND'}")

