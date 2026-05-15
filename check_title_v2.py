from docx import Document
from docx.oxml.ns import qn

doc = Document('/home/ubuntu/rencipe/docs/CPSC597 Project Report Draft (2).docx')
# Look at the first table if it exists
if doc.tables:
    first_table = doc.tables[0]
    for row in first_table.rows:
        for cell in row.cells:
            print(f"Cell text: {cell.text.strip()}")
            # Check for font name or size in the cell's paragraphs
            for p in cell.paragraphs:
                for run in p.runs:
                    if run.font.name:
                        print(f"Font name: {run.font.name}")
                    if run.font.size:
                        print(f"Font size: {run.font.size.pt}")

